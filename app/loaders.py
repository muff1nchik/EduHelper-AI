"""Извлекает текст из поддерживаемых учебных файлов."""

from abc import ABC, abstractmethod
from pathlib import Path

import fitz
from docx import Document

from app.text_utils import normalize_document_text


class BaseLoader(ABC):
    """Задаёт общий интерфейс для загрузки документов."""

    @abstractmethod
    def load_text(self, file_path: str) -> str:
        """Извлекает текст из файла."""

    def _ensure_not_empty(self, text: str) -> str:
        """Нормализует текст и проверяет, что он не пустой."""
        text = normalize_document_text(text)
        if not text or not text.strip():
            raise ValueError("Документ пустой или не содержит текстового слоя.")
        return text


class TextLoader(BaseLoader):
    """Загружает текстовые файлы TXT и Markdown."""

    def load_text(self, file_path: str) -> str:
        """Читает текстовый файл в кодировке UTF-8."""
        try:
            text = Path(file_path).read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError(
                "Не удалось прочитать текстовый файл в кодировке UTF-8."
            ) from exc
        return self._ensure_not_empty(text)


class PdfLoader(BaseLoader):
    """Загружает текст из PDF без OCR."""

    def load_text(self, file_path: str) -> str:
        """Извлекает текст из всех страниц PDF."""
        try:
            with fitz.open(file_path) as document:
                text = "\n\n".join(page.get_text("text") for page in document)
        except Exception as exc:
            raise ValueError("Не удалось прочитать PDF-файл.") from exc
        return self._ensure_not_empty(text)


class DocxLoader(BaseLoader):
    """Загружает текст из DOCX-документа."""

    def load_text(self, file_path: str) -> str:
        """Извлекает абзацы и таблицы из DOCX."""
        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError("Не удалось прочитать DOCX-файл.") from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text and cell.text.strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))

        return self._ensure_not_empty("\n".join(parts))


def get_loader(file_path: str) -> BaseLoader:
    """Выбирает загрузчик по расширению файла."""
    extension = Path(file_path).suffix.lower()
    if extension in {".txt", ".md"}:
        return TextLoader()
    if extension == ".pdf":
        return PdfLoader()
    if extension == ".docx":
        return DocxLoader()
    raise ValueError("Поддерживаются только файлы PDF, TXT, MD и DOCX.")
